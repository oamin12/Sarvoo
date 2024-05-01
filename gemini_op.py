import pathlib
import textwrap

import google.generativeai as genai

from IPython.display import display
from IPython.display import Markdown
from docx import Document
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from PIL import Image

def to_markdown(text):
  text = text.replace('•', '  *')
  return Markdown(textwrap.indent(text, '> ', predicate=lambda _: True))

# Used to securely store your API key
#from google.colab import userdata

GOOGLE_API_KEY="AIzaSyDU-8s2RdcFNrQ5Ca3IxVoBG8G1diDtryQ"

genai.configure(api_key=GOOGLE_API_KEY)

for m in genai.list_models():
  if 'generateContent' in m.supported_generation_methods:
    print(m.name)

model = genai.GenerativeModel('gemini-pro')

ai_propmt = '''Given the job description, CV, and interview transcript provided below, evaluate the candidate's suitability for the position. Provide a recommendation on whether to hire the candidate and summarize the key reasons for your decision. give a clear decision whether to hire the candidate or not'''
outputformat = '''I want to generate a consistent report format for the answer. The report should include the following sections:

Title: Hiring Recommendation Report for [Candidate Name]

1- Recommendation: A clear decision on whether to hire the candidate or not, based on the evaluation of the job description, CV, and interview transcript.



2- you will give a *Candidate Score* (out of 5; 5 is excellent, 4 is very good, 3 is good, 2 is fair, 1 is weak)  for each of the following criteria and provide a brief explanation for each criteria and a *reason* for the given score. The criteria are as follows:
1-Relevant Skills and Qualifications: Interviewers assess whether the interviewee possesses the necessary skills, experience, and qualifications to perform the job effectively. This includes technical skills, soft skills, and any industry-specific knowledge required for the role.
2-Cultural Fit: Employers look for candidates whose values, work ethic, and personality align with the company culture. They assess how well the interviewee will integrate into the team and contribute positively to the work environment.
3-Communication Skills: Effective communication is crucial in most roles. Interviewers evaluate how well the interviewee communicates ideas, listens to others, and articulates thoughts clearly and professionally.
4-Problem-Solving Abilities: Interviewers often present hypothetical scenarios or real-world problems to assess the interviewee's ability to think critically, analyze situations, and propose viable solutions. They look for candidates who demonstrate creativity, resourcefulness, and a proactive approach to problem-solving.
5-Adaptability and Flexibility: In today's dynamic work environments, adaptability is highly valued. Interviewers assess whether the interviewee can adapt to changing circumstances, learn new skills quickly, and thrive in diverse situations.

if you can't find a score for a criterion, you can leave it empty. and write "not applicable".

After writing the scores above, write them all together again in one line just like this: *Scores: num1, num2, num3, num4, num5*.

3- Summary of Reasons: A concise summary of the key reasons for the recommendation, highlighting the candidate's strengths and areas for improvement based on the evaluation of the job description, CV, and interview transcript.

put the first and second sections in the first page in a list format. and the third section in a paragraph format in a second page.

the page number should be the first thing at the top of each page.


'''

def generate_report(Job_description, CV, transcript):
    response = model.generate_content(ai_propmt + Job_description + CV + transcript + outputformat)

    display(to_markdown(response.text))
    text = response.text
    text = text.replace('**', '')
    lines = text.split('\n')

    page1 = []
    page2 = []
    page_number = 0
    for line in lines:
        if line.find('Page 1')!=-1:
            page_number = 1
            continue
        if line.find('Page 2')!=-1:
            page_number = 2
            continue
        if page_number == 1:
            page1.append(line)
        if page_number == 2:
            page2.append(line)
        if line.find('Scores:')!=-1:
            scores = line[line.find('Scores:')+7:].split(',')
            for i in range(len(scores)):
                #check if the score is not a number
                if scores[i].strip() == 'not applicable':
                    scores[i] = 0
            scores = [int(score.strip()) for score in scores]
            print(scores,"SSCORESS")
            scores = pd.Series(scores)
            #spider plot 
            fig = go.Figure()
            fig.add_trace(go.Scatterpolar(
                r=scores,
                theta=['Relevant Skills and Qualifications', 'Cultural Fit', 'Communication Skills', 'Problem-Solving Abilities', 'Adaptability and Flexibility'],
                fill='toself',
                name='Scores'
            ))
            fig.update_layout(
                polar=dict(
                    radialaxis=dict(
                        visible=True,
                        range=[0, 5]
                    )),
                showlegend=False
            )
            fig.show()
            fig.write_image("scores.png")

    page1 = '\n'.join(page1)
    page2 = '\n'.join(page2)

    print(page1)
    print(page2)

    #generate a word document
    doc = Document()
    doc.add_heading(lines[0], level=1)
    doc.add_heading("Recommendation and Evaluation", level=2)
    doc.add_paragraph(page1)
    doc.add_page_break()
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(page2)
    #resize the image
    image = Image.open("scores.png")
    new_image = image.resize((400, 400))
    new_image.save("scores.png")
    doc.add_picture("scores.png")
    doc.save('report.docx')

    return response








        
